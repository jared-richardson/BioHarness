#!/usr/bin/env python3
"""Export the manuscript markdown draft to a Word document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_ITEM_RE = re.compile(r"^\d+\.\s+(.*)$")
UNORDERED_ITEM_RE = re.compile(r"^-\s+(.*)$")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)$")
CAPTION_RE = re.compile(r"^\*\*((?:Figure|Table)\s+\d+\..*)\*\*$")
TABLE_SEPARATOR_RE = re.compile(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$")
INLINE_TOKEN_RE = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))",
)
LINK_RE = re.compile(r"^\[(.+?)\]\((.+?)\)$")


def _configure_document(document: Document) -> None:
    """Apply document-level styling."""
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    document.styles["Title"].font.name = "Times New Roman"
    document.styles["Title"].font.size = Pt(20)
    document.styles["Title"].font.bold = True

    for style_name, size in (
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def _resolve_link_text(label: str, target: str) -> str:
    """Render markdown links as readable text in the Word document."""
    if target.startswith("http://") or target.startswith("https://"):
        return f"{label} ({target})"
    return label


def _add_inline_markdown(paragraph, text: str) -> None:
    """Render simple inline markdown into a paragraph."""
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        else:
            link_match = LINK_RE.match(token)
            if link_match:
                label, target = link_match.groups()
                rendered = _resolve_link_text(label, target)
                run = paragraph.add_run(rendered)
                if target.startswith("http://") or target.startswith("https://"):
                    run.underline = True
            else:
                paragraph.add_run(token)
        position = match.end()

    if position < len(text):
        paragraph.add_run(text[position:])


def _is_table_start(lines: list[str], index: int) -> bool:
    """Return True when the current line starts a markdown table."""
    if index + 1 >= len(lines):
        return False
    return lines[index].startswith("|") and bool(TABLE_SEPARATOR_RE.match(lines[index + 1]))


def _split_table_row(line: str) -> list[str]:
    """Split a markdown table row into cells."""
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _add_table(document: Document, rows: list[list[str]]) -> None:
    """Add a markdown table to the Word document."""
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, cell_text in enumerate(row):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            _add_inline_markdown(paragraph, cell_text)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def _is_horizontal_rule(line: str) -> bool:
    """Return True when a line is a horizontal rule marker."""
    stripped = line.strip()
    return stripped in {"---", "***"}


def _add_caption(document: Document, text: str) -> None:
    """Add a centered italic caption."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def _add_image(document: Document, manuscript_dir: Path, target: str) -> None:
    """Embed an image referenced from the markdown manuscript."""
    image_path = (manuscript_dir / target).resolve()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))


def _collect_paragraph(lines: list[str], index: int) -> tuple[str, int]:
    """Collect a paragraph block until the next structural marker."""
    collected: list[str] = []
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            break
        if (
            HEADING_RE.match(line)
            or IMAGE_RE.match(line.strip())
            or CAPTION_RE.match(line.strip())
            or _is_horizontal_rule(line)
            or ORDERED_ITEM_RE.match(line)
            or UNORDERED_ITEM_RE.match(line)
            or _is_table_start(lines, index)
        ):
            break
        collected.append(line.strip())
        index += 1
    return " ".join(collected), index


def _add_list(document: Document, items: list[str], ordered: bool) -> None:
    """Add a bulleted or numbered list."""
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        paragraph = document.add_paragraph(style=style)
        _add_inline_markdown(paragraph, item)


def export_manuscript(markdown_path: Path, output_path: Path) -> None:
    """Export the markdown manuscript to a Word document."""
    document = Document()
    _configure_document(document)
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    manuscript_dir = markdown_path.parent

    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if _is_horizontal_rule(stripped):
            index += 1
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if first_heading and level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_inline_markdown(paragraph, text)
                first_heading = False
            else:
                paragraph = document.add_paragraph(style=f"Heading {min(level, 3)}")
                _add_inline_markdown(paragraph, text)
            index += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            _, target = image_match.groups()
            _add_image(document, manuscript_dir, target)
            index += 1
            continue

        caption_match = CAPTION_RE.match(stripped)
        if caption_match:
            _add_caption(document, caption_match.group(1))
            index += 1
            continue

        if _is_table_start(lines, index):
            rows: list[list[str]] = []
            rows.append(_split_table_row(lines[index]))
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(_split_table_row(lines[index]))
                index += 1
            _add_table(document, rows)
            continue

        if ORDERED_ITEM_RE.match(line):
            items: list[str] = []
            while index < len(lines):
                match = ORDERED_ITEM_RE.match(lines[index])
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            _add_list(document, items, ordered=True)
            continue

        if UNORDERED_ITEM_RE.match(line):
            items = []
            while index < len(lines):
                match = UNORDERED_ITEM_RE.match(lines[index])
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            _add_list(document, items, ordered=False)
            continue

        paragraph_text, index = _collect_paragraph(lines, index)
        if paragraph_text:
            paragraph = document.add_paragraph()
            if paragraph_text.startswith("**Author:**"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_markdown(paragraph, paragraph_text)
        else:
            index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    """Run the manuscript export from the repository defaults."""
    repo_root = Path(__file__).resolve().parents[1]
    markdown_path = repo_root / "docs" / "manuscript.md"
    output_path = repo_root / "docs" / "manuscript.docx"
    export_manuscript(markdown_path, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
